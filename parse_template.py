#!/usr/bin/env python3
"""Script to parse Character Template.docx into JSON structure."""

import json
from pathlib import Path
from docx import Document


def parse_docx_to_json(docx_path: str) -> dict:
    """Parse a DOCX file into a JSON structure based on headers and fields.
    
    Args:
        docx_path: Path to the DOCX file
        
    Returns:
        Dictionary with structure based on document headers and fields
    """
    doc = Document(docx_path)
    result = {}
    current_section = None
    
    print(f"📖 Parsing: {docx_path}")
    print(f"📋 Total paragraphs: {len(doc.paragraphs)}\n")
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        
        # Get paragraph style
        style = para.style.name if para.style else "Normal"
        
        # Detect heading level from style
        level = _get_heading_level(style)
        
        if level == 1:  # Main section heading
            current_section = text
            result[current_section] = []
            print(f"📌 SECTION: {current_section}")
            
        elif current_section and _is_field_like(text):
            # This is a field under the current section
            # Store as-is, removing trailing colon for consistency
            field_name = text.rstrip(":").strip()
            result[current_section].append(field_name)
            print(f"  • {field_name}")
    
    print()
    return result


def _get_heading_level(style_name: str) -> int:
    """Extract heading level from style name.
    
    Args:
        style_name: Style name from DOCX
        
    Returns:
        Heading level (1-3) or 0 for non-headings
    """
    if "Heading 1" in style_name:
        return 1
    elif "Heading 2" in style_name:
        return 2
    elif "Heading 3" in style_name:
        return 3
    return 0


def _is_field_like(text: str) -> bool:
    """Check if text looks like a field (label with colon or placeholder).
    
    Args:
        text: Text to check
        
    Returns:
        True if looks like a field
    """
    # Contains placeholder markers
    if "{{" in text or "**" in text:
        return True
    # Ends with colon
    if text.endswith(":"):
        return True
    return False


if __name__ == "__main__":
    template_path = "Character Template.docx"
    
    if not Path(template_path).exists():
        print(f"❌ File not found: {template_path}")
        exit(1)
    
    # Parse the document
    structure = parse_docx_to_json(template_path)
    
    # Output JSON
    print("\n" + "="*60)
    print("📊 PARSED JSON STRUCTURE:")
    print("="*60 + "\n")
    print(json.dumps(structure, indent=2))
    
    # Save to file
    output_path = "character_template_structure.json"
    with open(output_path, 'w') as f:
        json.dump(structure, f, indent=2)
    print(f"\n✅ Saved to: {output_path}")
